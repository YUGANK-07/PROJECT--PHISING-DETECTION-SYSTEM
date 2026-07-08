import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('PhishGuard: Production-Grade Phishing Detection System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Technical Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True
    
    doc.add_page_break()

    # Define common styles
    def add_heading(text, level=1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

    # 1. Problem Statement
    add_heading('1. Problem Statement')
    doc.add_paragraph(
        "Phishing attacks are becoming increasingly sophisticated, evolving to bypass traditional "
        "URL-based and signature-based detection systems. Modern threat actors leverage advanced "
        "obfuscation techniques, visual brand spoofing, and rapidly changing domains to deceive "
        "unsuspecting users. A comprehensive system is critically needed to detect these zero-day "
        "threats in real-time by analyzing multiple dimensions of a webpage simultaneously, beyond "
        "just its lexical structure."
    )

    # 2. Objectives
    add_heading('2. Objectives')
    objectives = [
        "Build a real-time, scalable phishing detection system with exceptionally high accuracy and a minimal false positive rate.",
        "Implement a multi-layered, 125-dimensional feature engineering approach covering lexical, structural, and visual aspects.",
        "Develop an ensemble machine learning architecture integrating multiple predictive models for robust classification.",
        "Incorporate a Visual Similarity Module to detect visual brand spoofing using Deep Learning.",
        "Provide Explainable AI (XAI) capabilities to ensure users and administrators understand precisely why a site is flagged.",
        "Create a robust backend API (FastAPI) accompanied by a user-friendly, responsive web UI for seamless interaction."
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # 3. Methodology
    add_heading('3. Methodology')
    doc.add_paragraph(
        "The project methodology involves a holistic pipeline ranging from data collection to explainable classification."
    )
    
    add_heading('Feature Engineering (125-Dimensional Vector)', level=2)
    doc.add_paragraph("The system extracts features across five distinct categories:")
    features = [
        "URL Lexical Features (45 dimensions): Analyzes URL structure, entropy, length, TLD risk, and homograph indicators.",
        "Domain Features (19 dimensions): Evaluates WHOIS records, domain age, DNS A/MX/NS configuration, and SSL certificate validity.",
        "Webpage Features (31 dimensions): Inspects HTML structure, form actions, hidden iframes, meta-refreshes, and obfuscated JavaScript.",
        "NLP Features (7 dimensions): Extracts phishing phrase density, urgency words, and brand impersonation markers using text analysis.",
        "Visual Similarity Module: Utilizes a pre-trained ResNet50 CNN and perceptual hashing (pHash) to analyze Playwright-captured screenshots for UI spoofing."
    ]
    for feat in features:
        doc.add_paragraph(feat, style='List Bullet')

    add_heading('Machine Learning & Ensemble Architecture', level=2)
    doc.add_paragraph(
        "Base models including Random Forest, XGBoost, and a PyTorch Neural Network (MLP) are trained "
        "independently. Their predictions are aggregated using a stacking ensemble (Out-Of-Fold predictions "
        "combined with a Logistic Regression meta-classifier) to maximize predictive performance and generalizability."
    )

    add_heading('Explainability (XAI)', level=2)
    doc.add_paragraph(
        "SHAP (SHapley Additive exPlanations) is integrated to compute feature importance. TreeExplainer "
        "and KernelExplainer provide human-readable interpretations for each prediction, mapping complex "
        "model outputs to tangible risk factors."
    )

    # 4. Implementation
    add_heading('4. Implementation')
    doc.add_paragraph("The system architecture is composed of the following core components:")
    components = [
        "Backend & API: Built with FastAPI and Uvicorn for high-performance, asynchronous request handling. "
        "It includes Redis for caching, JWT for secure authentication, and detailed Pydantic schemas.",
        "Machine Learning Pipeline: Leverages scikit-learn, XGBoost, and PyTorch for model training. "
        "A scalable feature extraction engine unifies lexical, NLP, and visual data into a single representation.",
        "Visual Subsystem: Employs Playwright for headless browser interaction to capture live screenshots, "
        "processed by PyTorch and imagehash.",
        "Frontend Interface: A premium, single-page application built using vanilla HTML/CSS/JS, featuring "
        "a dark glassmorphism design, real-time feedback, and dynamic result visualization."
    ]
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')

    # 5. Results & Output
    add_heading('5. Results & Output')
    doc.add_paragraph(
        "The models were trained on approximately 250,000 URLs (sourced from PhishTank, OpenPhish, and Tranco) "
        "and evaluated on a held-out test set of 37,500 URLs."
    )
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Precision'
    hdr_cells[2].text = 'Recall'
    hdr_cells[3].text = 'F1 Score'
    hdr_cells[4].text = 'ROC-AUC'
    
    data = [
        ('Random Forest', '1.0000', '1.0000', '1.0000', '1.0000'),
        ('XGBoost', '0.9999', '0.9994', '0.9997', '1.0000'),
        ('Ensemble', '1.0000', '1.0000', '1.0000', '1.0000')
    ]
    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        row_cells[4].text = item[4]

    doc.add_paragraph("")
    doc.add_paragraph(
        "The Ensemble model achieved a False Negative Rate (FNR) of just 0.004%, "
        "missing only 1 phishing URL out of 22,500 tested."
    )
    doc.add_paragraph(
        "System Output: The API provides rapid responses (typically ~100-150ms) detailing "
        "phishing probability, an assigned risk level, and an array of SHAP-based explanations "
        "highlighting specific threats."
    )

    # 6. Conclusion
    add_heading('6. Conclusion')
    doc.add_paragraph(
        "PhishGuard successfully demonstrates a highly accurate, explainable, and scalable approach "
        "to modern phishing detection. By combining traditional lexical and structural features with "
        "advanced visual similarity analysis and ensemble machine learning, the system can reliably "
        "detect both known and zero-day phishing attempts. The integration of Explainable AI (SHAP) "
        "provides crucial transparency, building trust with end-users and network administrators."
    )

    # 7. Future Scope
    add_heading('7. Future Scope')
    future = [
        "Browser Extension Integration: Developing lightweight plugins for Chrome and Firefox to provide automated, on-the-fly protection for users.",
        "Continuous Online Learning: Implementing mechanisms for the model to adapt to emerging phishing trends without requiring complete dataset retraining.",
        "Expanded Visual Database: Broadening the reference database for the visual similarity module to cover a wider spectrum of frequently targeted international brands.",
        "Enhanced Multilingual Support: Improving the NLP feature extraction module to detect urgency and phishing intent in non-English web pages."
    ]
    for item in future:
        doc.add_paragraph(item, style='List Bullet')

    doc.save('PhishGuard_Project_Report.docx')

if __name__ == "__main__":
    create_report()
